import os
import time
import streamlit as st
import google.generativeai as genai
import pypdf
import docx

# --- 0. PAGE & THEME -------------------------------------------------------
st.set_page_config(
    page_title="Sportfysioloog AI",
    page_icon="🚴‍♂️",
    layout="wide",
)

# Global look & feel (strak, blauwe achtergrond)
st.markdown(
    r"""
    <style>
      :root {
        --bg: linear-gradient(145deg, #061428 0%, #0a2f4a 55%, #0c4c63 100%);
        --card: rgba(255, 255, 255, 0.05);
        --border: rgba(255, 255, 255, 0.12);
        --text: #eef5ff;
        --muted: #b9c6da;
        --accent: #7be0ff;
      }
      body, .stApp { background: var(--bg); color: var(--text); font-family: "Inter", -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; }
      .block-container { padding: 1.5rem 2rem 2.5rem; max-width: 1100px; }
      .hero, .upload-card, .chat-card {
        background: var(--card);
        border: 1px solid var(--border);
        border-radius: 12px;
        padding: 1rem 1.1rem;
        box-shadow: 0 20px 70px rgba(0,0,0,0.25);
      }
      .hero h1 { margin: 0 0 .35rem; font-weight: 700; color: var(--text); font-size: 32px; white-space: nowrap; }
      @media (max-width: 760px) { .hero h1 { white-space: normal; font-size: 26px; } }
      .hero p { color: var(--muted); font-size: 0.95rem; margin: 0.1rem 0 0; }
      .cta-btn button { width:100%; background: var(--text); color: #041021; border:none; border-radius:12px; padding: 0.65rem 0; }
      .stExpander, .stFileUploader { color: var(--text); }
      .stChatMessage { background: transparent; }
      .stMarkdown p { color: var(--text); }
      .chat-card .stChatMessage[data-testid="stChatMessage"] div { color: var(--text); }
      .bike-loader { display:flex; gap:6px; font-size:24px; margin: 4px 0 0; color: var(--muted); }
      .bike-loader div { animation: ride 0.9s ease-in-out infinite; }
      .bike-loader div:nth-child(2) { animation-delay: .15s; }
      .bike-loader div:nth-child(3) { animation-delay: .3s; }
      @keyframes ride { 0% { transform: translateX(0px); } 50% { transform: translateX(10px); } 100% { transform: translateX(0px); } }
      /* input bar in same blue zone */
      div[data-testid="stChatInput"] { background: transparent; }
      div[data-testid="stChatInput"] textarea { background: rgba(255,255,255,0.06); border: 1px solid var(--border); color: var(--text); border-radius: 14px; }
      /* file drop smaller & round */
      .stFileUploader div[data-testid="stFileDropzone"] { padding: 0.55rem 0.8rem; border-radius: 12px; }
      .stFileUploader button { border-radius: 10px !important; padding: 0.4rem 0.85rem !important; }
      /* containers narrower for upload row on wide screens */
      @media (min-width: 1024px) { .upload-row { max-width: 820px; } }
      header { margin-bottom: 0; }
    </style>
    """,
    unsafe_allow_html=True,
)

# --- 1. LOGO & MODEL -------------------------------------------------------
LOGO_PATH = "1.png"  # pas eventueel aan als je logo-bestand anders heet
MODEL_NAME = "gemini-2.5-flash"  # expliciet 2.5 gebruiken

# --- 2. CONFIGURATIE & API ------------------------------------------------
try:
    if "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"].strip()
        genai.configure(api_key=api_key)
    else:
        st.error("Geen API Key gevonden.")
        st.stop()
except Exception as e:
    st.error(f"Error: {e}")
    st.stop()

# --- 3. KENNIS LADEN (PDF & DOCX) -----------------------------------------
@st.cache_resource(show_spinner=False)
def load_all_knowledge():
    """Zoekt automatisch naar alle PDF en DOCX bestanden en leest ze."""
    combined_text = ""
    for filename in os.listdir("."):
        try:
            if filename.lower().endswith(".pdf"):
                reader = pypdf.PdfReader(filename)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        combined_text += text + "\n"
            elif filename.lower().endswith(".docx"):
                doc = docx.Document(filename)
                for para in doc.paragraphs:
                    combined_text += para.text + "\n"
        except Exception as e:
            print(f"Kon bestand {filename} niet lezen: {e}")
    return combined_text

knowledge_base = load_all_knowledge()

# --- 4. AI INSTRUCTIES ----------------------------------------------------
SYSTEM_PROMPT = f"""
ROL: Je bent een expert sportfysioloog van SportMetrics.

BRONMATERIAAL:
Je hebt toegang tot specifieke literatuur over trainingsleer (zie hieronder).
Gebruik DEZE INFORMATIE als de absolute waarheid.

=== START LITERATUUR ===
{knowledge_base}
=== EINDE LITERATUUR ===

BELANGRIJKE REGELS:
1. SportMetrics doet GEEN lactaatmetingen (prikken), alleen ademgasanalyse.
2. Gebruik de principes (zoals Seiler zones) zoals beschreven in de geüploade literatuur.
3. Wees praktisch, enthousiast en gebruik bulletpoints.
4. Geen medisch advies.
5. Geef altijd een props aan de persoon voor de test en bedankt dat hij of zij dit bij SportMetrics heeft gedaan.
"""

try:
    model = genai.GenerativeModel(
        model_name=MODEL_NAME,
        system_instruction=SYSTEM_PROMPT,
    )
except Exception as e:
    st.error(f"Model fout ({MODEL_NAME}): {e}")

# --- 5. HERO --------------------------------------------------------------
hero = st.container()
with hero:
    col1, col2 = st.columns([1.5, 1])
    with col1:
        st.markdown("<h1>SportMetrics Wieler & Hardloop Coach</h1>", unsafe_allow_html=True)
        st.markdown(
            "<p>Ik ben een AI-model: ik geef zo goed mogelijk advies op basis van de beste literatuur, no-nonsense, praktisch en to the point. Geen medisch advies; alles dankzij jouw inspanningstest bij SportMetrics.</p>",
            unsafe_allow_html=True,
        )
        st.caption("Geen poespas, alleen data → advies.")
    with col2:
        if os.path.exists(LOGO_PATH):
            st.image(LOGO_PATH, width=200)
        else:
            st.info("Plaats je logo als '1.png' in de map om het hier te tonen.")

# --- 6. UPLOAD ------------------------------------------------------------
upload_card = st.container()
with upload_card:
    st.markdown("<div class='upload-card upload-row'>", unsafe_allow_html=True)
    col_u1, col_u2 = st.columns([1.4, 1])
    with col_u1:
        st.subheader("Upload rapport")
        uploaded_file = st.file_uploader("Kies PDF of DOCX", type=["pdf", "docx"], label_visibility="collapsed")
    with col_u2:
        st.markdown("<div class='cta-btn'>", unsafe_allow_html=True)
        st.button("Analyseer", use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)
        st.caption("Tip: zeg 'Maak mijn zones'.")
    st.markdown("</div>", unsafe_allow_html=True)

if uploaded_file is not None:
    try:
        client_pdf_text = ""
        if uploaded_file.name.lower().endswith(".pdf"):
            reader = pypdf.PdfReader(uploaded_file)
            for page in reader.pages:
                client_pdf_text += page.extract_text() + "\n"
        elif uploaded_file.name.lower().endswith(".docx"):
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                client_pdf_text += para.text + "\n"
        st.session_state["last_uploaded_text"] = client_pdf_text
        st.toast("Rapport ontvangen! Typ je vraag beneden.", icon="✅")
    except Exception as e:
        st.error(f"Fout bij lezen rapport: {e}")

# --- 7. CHAT HISTORY ------------------------------------------------------
if "messages" not in st.session_state:
    st.session_state.messages = []
    intro = (
        "Hoi! Ik geef antwoord op basis van mijn AI-kennis en de best beschikbare literatuur over trainingsleer.\n\n"
        "Upload je testresultaten of stel direct een vraag!"
    )
    st.session_state.messages.append({"role": "assistant", "content": intro})

chat_box = st.container()
with chat_box:
    st.markdown("<div class='chat-card'>", unsafe_allow_html=True)
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    st.markdown("</div>", unsafe_allow_html=True)

# --- 8. CHAT INPUT --------------------------------------------------------
prompt = st.chat_input("Stel je vraag of zeg 'Maak mijn zones'...")

if prompt:
    extra_context = ""
    if "last_uploaded_text" in st.session_state:
        extra_context = f"\n\nHIER IS HET RAPPORT VAN DE KLANT:\n{st.session_state['last_uploaded_text']}\n\n"
        del st.session_state["last_uploaded_text"]

    full_prompt_for_ai = prompt + extra_context

    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    try:
        with st.chat_message("assistant"):
            bike_placeholder = st.empty()
            bike_placeholder.markdown(
                """
                <div class="bike-loader">
                  <div>🚴‍♀️</div><div>🚴‍♂️</div><div>🚴‍♀️</div>
                </div>
                <p style="color:var(--muted); margin-top:2px;">Antwoord wordt geladen...</p>
                """,
                unsafe_allow_html=True,
            )
            with st.spinner("🚴‍♂️🚴‍♀️ bezig met jouw advies..."):
                response = model.generate_content(full_prompt_for_ai)
            bike_placeholder.markdown(response.text)
            st.session_state.messages.append({"role": "assistant", "content": response.text})
    except Exception as e:
        st.error(f"De AI reageert niet: {e}")
